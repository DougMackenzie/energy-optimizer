"""
Enhanced Report Builder with Real Google Sheets Data & AI Analysis
Generates comprehensive Word reports with actual optimization results, charts, and AI-generated insights
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import os
from typing import List, Dict, Optional
import pandas as pd

# Import our new modules
try:
    from app.utils.portfolio_data import (
        load_all_site_results,
        get_portfolio_summary,
        get_site_optimization_history,
        get_equipment_summary,
        fetch_site_results_from_sheets
    )
    from app.utils.enhanced_report_charts import (
        create_lcoe_comparison_chart,
        create_capex_breakdown_chart,
        create_15year_energy_stack_chart,
        create_flexibility_impact_chart,
        create_site_map_image
    )
    from app.utils.report_charts import (
        create_8760_dispatch_chart,
        create_deployment_timeline_chart
    )
    from app.utils.gemini_client import GeminiReportClient
    ENHANCED_FEATURES_AVAILABLE = True
except ImportError as e:
    print(f"Enhanced features not available: {e}")
    ENHANCED_FEATURES_AVAILABLE = False


def generate_enhanced_word_report(site_selection: List[str], content_options: Dict, 
                                   sites_list: List[Dict] = None, 
                                   use_ai: bool = True) -> bytes:
    """
    Generate comprehensive Word document report with real data from Google Sheets
    
    Args:
        site_selection: List of selected site names
        content_options: Dict of content sections to include (True/False)
        sites_list: Full list of sites (optional, will fetch from Sheets if None)
        use_ai: Whether to use Gemini AI for text generation
    
    Returns:
        Word document as bytes
    """
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Energy Optimization Report"
    doc.core_properties.author = "Antigravity Energy Optimizer"
    doc.core_properties.created = datetime.now()
    
    print(f"\n=== ENHANCED REPORT GENERATION DEBUG ===")
    print(f"Site selection: {site_selection}")
    print(f"Enhanced features available: {ENHANCED_FEATURES_AVAILABLE}")
    print(f"Use AI: {use_ai}")
    
    # Initialize AI client if enabled
    ai_client = None
    if use_ai and ENHANCED_FEATURES_AVAILABLE:
        try:
            ai_client = GeminiReportClient()
            print(f"✓ AI client initialized: {ai_client.model_name}")
        except Exception as e:
            print(f"✗ AI client unavailable: {e}")
    
    # Fetch site data from Google Sheets
    site_results = load_all_site_results() if ENHANCED_FEATURES_AVAILABLE else []
    print(f"Loaded {len(site_results)} sites from Google Sheets")
    
    # Filter by selection
    if "📊 Entire Portfolio" not in site_selection:
        site_results = [s for s in site_results if s.get('site_name') in site_selection]
    
    print(f"After filtering: {len(site_results)} sites")
    for sr in site_results:
        print(f"  - {sr.get('site_name')}: LCOE ${sr.get('lcoe', 0):.1f}/MWh, Stage: {sr.get('stage')}")
    
    # =============================================================================
    # Title Page
    # =============================================================================
    if content_options.get('include_exec_overview', True):
        print("Adding title page...")
        add_enhanced_title_page(doc, site_selection, site_results)
        doc.add_page_break()
    
    # =============================================================================
    # Executive Summary with AI
    # =============================================================================
    if content_options.get('include_exec_metrics', True):
        print("Adding executive summary...")
        add_enhanced_executive_summary(doc, site_results, ai_client)
        doc.add_page_break()
    
    # =============================================================================
    # Financial Analysis with Charts
    # =============================================================================
    if content_options.get('include_cash_flow', True) or content_options.get('include_npv_irr', True):
        print("Adding financial analysis...")
        add_enhanced_financial_analysis(doc, site_results, content_options, ai_client)
        doc.add_page_break()
    
    # =============================================================================
    # Technical Analysis with Equipment Details
    # =============================================================================
    if content_options.get('include_equipment', True) or content_options.get('include_optimization', True):
        print("Adding technical analysis...")
        add_enhanced_technical_analysis(doc, site_results, content_options, ai_client)
        doc.add_page_break()
    
    # =============================================================================
    # Load Profile & Dispatch (8760 Sample Week)
    # =============================================================================
    if content_options.get('include_load_profile', True):
        print("Adding 8760 dispatch section...")
        add_8760_dispatch_section(doc, site_results)
        doc.add_page_break()
    
    # =============================================================================
    # 15-Year Energy Stack
    # =============================================================================
    if content_options.get('include_stage_progression', True):
        print("Adding 15-year energy stack...")
        add_15year_energy_stack(doc, site_results)
        doc.add_page_break()
    
    # =============================================================================
    # Site Maps (if GeoJSON available)
    # =============================================================================
    if content_options.get('include_location_map', True):
        print("Adding site maps...")
        add_site_maps_section(doc, site_results)
    
    print("=== REPORT GENERATION COMPLETE ===\n")
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def add_enhanced_title_page(doc: Document, site_selection: List[str], site_results: List[Dict]):
    """Add title page with real portfolio metrics"""
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ENERGY OPTIMIZATION REPORT")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph()
    
    # Subtitle - Portfolio or Site-specific
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if "📊 Entire Portfolio" in site_selection or len(site_results) > 1:
        portfolio_summary = get_portfolio_summary(site_results) if ENHANCED_FEATURES_AVAILABLE else {}
        subtitle_text = f"Portfolio Analysis - {portfolio_summary.get('num_sites', len(site_results))} Sites"
        run = subtitle.add_run(subtitle_text)
        run.font.size = Pt(18)
        
        # Add key portfolio metrics
        doc.add_paragraph()
        metrics_para = doc.add_paragraph()
        metrics_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metrics_text = f"Total Capacity: {portfolio_summary.get('total_capacity_mw', 0):.0f} MW | "
        metrics_text += f"Weighted LCOE: ${portfolio_summary.get('weighted_lcoe', 0):.1f}/MWh"
        run = metrics_para.add_run(metrics_text)
        run.font.size = Pt(12)
        run.font.italic = True
    else:
        site_name = site_results[0].get('site_name', 'Unknown Site') if site_results else site_selection[0]
        run = subtitle.add_run(f"{site_name} - Detailed Analysis")
        run.font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)


def add_enhanced_executive_summary(doc: Document, site_results: List[Dict], ai_client=None):
    """Add executive summary with AI-generated insights and real data"""
    
    # Section heading
    heading = doc.add_heading('Executive Summary', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    # Calculate portfolio metrics
    portfolio = get_portfolio_summary(site_results) if ENHANCED_FEATURES_AVAILABLE else {}
    
    # AI-generated summary (if available)
    if ai_client and portfolio:
        try:
            print(f"  Generating AI summary with {ai_client.model_name}...")
            ai_summary = ai_client.generate_executive_summary(portfolio)
            print(f"  ✓ AI summary generated ({len(ai_summary)} characters)")
            doc.add_paragraph(ai_summary)
        except Exception as e:
            print(f"  ✗ AI generation failed: {e}")
            # Fallback to template-based summary
            add_template_summary(doc, portfolio, site_results)
    else:
        print(f"  Using template summary (AI client: {ai_client is not None}, portfolio: {len(portfolio) > 0 if portfolio else False})")
        add_template_summary(doc, portfolio, site_results)
    
    doc.add_paragraph()
    
    # Key Metrics Table
    doc.add_paragraph("Portfolio Key Metrics:", style='Heading 2')
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    metrics_data = [
        ("Total Sites", f"{portfolio.get('num_sites', 0)}"),
        ("Total IT Capacity", f"{portfolio.get('total_capacity_mw', 0):.0f} MW"),
        ("Weighted Average LCOE", f"${portfolio.get('weighted_lcoe', 0):.1f}/MWh"),
        ("Total NPV", f"${portfolio.get('total_npv_m', 0):.1f}M"),
        ("Total CapEx Required", f"${portfolio.get('total_capex_m', 0):.1f}M"),
        ("Portfolio IRR (Est.)", f"{portfolio.get('portfolio_irr', 0):.1f}%")
    ]
    
    for i, (metric, value) in enumerate(metrics_data):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True


def add_template_summary(doc: Document, portfolio: Dict, site_results: List[Dict]):
    """Fallback template-based summary when AI is unavailable"""
    
    summary_text = (
        f"This report presents a comprehensive analysis of {portfolio.get('num_sites', 0)} energy optimization "
        f"sites with a combined IT capacity of {portfolio.get('total_capacity_mw', 0):.0f} MW. "
        f"The portfolio demonstrates a weighted average LCOE of ${portfolio.get('weighted_lcoe', 0):.1f}/MWh "
        f"with a total estimated NPV of ${portfolio.get('total_npv_m', 0):.1f}M."
    )
    doc.add_paragraph(summary_text)
    
    doc.add_paragraph(
        f"Total capital expenditure required is estimated at ${portfolio.get('total_capex_m', 0):.1f}M "
        f"with an estimated portfolio-level IRR of {portfolio.get('portfolio_irr', 0):.1f}%."
    )


def add_enhanced_financial_analysis(doc: Document, site_results: List[Dict], 
                                      content_options: Dict, ai_client=None):
    """Add financial analysis with comparative charts and AI insights"""
    
    # Section heading
    heading = doc.add_heading('Financial Analysis', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    # LCOE Comparison Chart
    if len(site_results) > 1 and content_options.get('include_lcoe_trend', True):
        doc.add_paragraph("LCOE Comparison Across Sites:", style='Heading 2')
        
        # Generate chart
        print(f"  Generating LCOE comparison chart...")
        try:
            chart_path = create_lcoe_comparison_chart(site_results)
            print(f"  Chart saved to: {chart_path}")
            if chart_path and os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.5))
                print(f"  ✓ LCOE chart embedded")
                doc.add_paragraph()  # Spacing
            else:
                print(f"  ✗ Chart path doesn't exist: {chart_path}")
        except Exception as e:
            print(f"  ✗ Failed to generate LCOE chart: {e}")
    
    # CapEx Breakdown (for each site or aggregate)
    if content_options.get('include_cash_flow', True):
        doc.add_paragraph("Capital Expenditure Breakdown:", style='Heading 2')
        
        for site in site_results[:3]:  # Limit to first 3 sites for brevity
            doc.add_paragraph(f"{site.get('site_name')}:", style='Heading 3')
            
            equipment = site.get('equipment', {})
            equipment_summary = get_equipment_summary(equipment) if ENHANCED_FEATURES_AVAILABLE else equipment
            
            # Generate CapEx chart
            chart_path = create_capex_breakdown_chart(equipment_summary)
            if chart_path and os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.5))
                
                # AI analysis of financial results
                if ai_client:
                    try:
                        site_financial_data = {
                            'site_name': site.get('site_name'),
                            'lcoe': site.get('lcoe', 0),
                            'npv_m': site.get('npv', 0) / 1_000_000,
                            'capex_m': equipment_summary.get('total_btm_mw', 0) * 1.5,  # Rough estimate
                            'capacity_mw': site.get('it_capacity_mw', 0)
                        }
                        ai_analysis = ai_client.analyze_financial_results(site_financial_data)
                        doc.add_paragraph(ai_analysis)
                    except:
                        pass
                
                doc.add_paragraph()  # Spacing


def add_enhanced_technical_analysis(doc: Document, site_results: List[Dict], 
                                      content_options: Dict, ai_client=None):
    """Add technical analysis with equipment specifications"""
    
    # Section heading
    heading = doc.add_heading('Technical Analysis', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    for site in site_results:
        doc.add_paragraph(f"{site.get('site_name')} - {site.get('stage', 'Detailed').capitalize()} Stage:", 
                          style='Heading 2')
        
        equipment = site.get('equipment', {})
        equipment_summary = get_equipment_summary(equipment) if ENHANCED_FEATURES_AVAILABLE else equipment
        
        # AI technical interpretation
        if ai_client and equipment:
            try:
                site_info = {
                    'site_name': site.get('site_name'),
                    'it_capacity_mw': site.get('it_capacity_mw', 0)
                }
                ai_tech_analysis = ai_client.interpret_technical_results(equipment_summary, site_info)
                doc.add_paragraph(ai_tech_analysis)
            except:
                pass
        
        # Equipment specs table
        doc.add_paragraph("Equipment Configuration:", style='Heading 3')
        
        equip_data = []
        if equipment_summary.get('recip_mw', 0) > 0:
            equip_data.append(('Reciprocating Engines', f"{equipment_summary['recip_mw']:.0f} MW"))
        if equipment_summary.get('turbine_mw', 0) > 0:
            equip_data.append(('Gas Turbines', f"{equipment_summary['turbine_mw']:.0f} MW"))
        if equipment_summary.get('bess_mwh', 0) > 0:
            equip_data.append(('Battery Storage', f"{equipment_summary['bess_mwh']:.0f} MWh"))
        if equipment_summary.get('solar_mw', 0) > 0:
            equip_data.append(('Solar PV', f"{equipment_summary['solar_mw']:.0f} MW DC"))
        if equipment_summary.get('grid_mw', 0) > 0:
            equip_data.append(('Grid Interconnect', f"{equipment_summary['grid_mw']:.0f} MW"))
        
        if equip_data:
            for equipment_type, capacity in equip_data:
                doc.add_paragraph(f"  • {equipment_type}: {capacity}", style='List Bullet')
        
        doc.add_paragraph()  # Spacing


def add_8760_dispatch_section(doc: Document, site_results: List[Dict]):
    """Add 8760 dispatch visualization (sample week)"""
    
    heading = doc.add_heading('Dispatch Profile - Sample Week', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph(
        "The following chart shows the hourly dispatch of generation resources over a representative "
        "week (168 hours). This visualization demonstrates how different energy sources are deployed "
        "to meet the data center load profile."
    )
    
    # Generate chart for first site (or aggregate for portfolio)
    if site_results:
        site = site_results[0]
        equipment = site.get('equipment', {})
        
        # Create equipment config format expected by chart function
        equipment_config = {
            'recip_engines': [{'capacity_mw': equipment.get('recip_mw', 0)}] if equipment.get('recip_mw', 0) > 0 else [],
            'gas_turbines': [{'capacity_mw': equipment.get('turbine_mw', 0)}] if equipment.get('turbine_mw', 0) > 0 else [],
            'bess': [{'power_mw': equipment.get('bess_mwh', 0) / 4}] if equipment.get('bess_mwh', 0) > 0 else [],  # 4-hour duration
            'solar_mw_dc': equipment.get('solar_mw', 0),
            'grid_import_mw': equipment.get('grid_mw', 0)
        }
        
        site_info = {
            'Total_Facility_MW': site.get('facility_mw', 200),
            'Load_Factor_Pct': 70
        }
        
        chart_path = create_8760_dispatch_chart(equipment_config, site_info)
        if chart_path and os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(6.5))


def add_15year_energy_stack(doc: Document, site_results: List[Dict]):
    """Add 15-year energy generation stack"""
    
    heading = doc.add_heading('15-Year Energy Generation Forecast', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph(
        "This section presents the projected annual energy generation by source over the 15-year analysis period, "
        "accounting for equipment degradation (solar: 0.5%/year, BESS: 2%/year)."
    )
    
    if site_results:
        site = site_results[0]
        equipment = site.get('equipment', {})
        equipment_summary = get_equipment_summary(equipment) if ENHANCED_FEATURES_AVAILABLE else equipment
        
        load_data = {
            'total_annual_gwh': site.get('it_capacity_mw', 200) * 8760 * 0.7 / 1000
        }
        
        chart_path = create_15year_energy_stack_chart(equipment_summary, load_data)
        if chart_path and os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(7))


def add_site_maps_section(doc: Document, site_results: List[Dict]):
    """Add site location maps with GeoJSON boundaries"""
    
    heading = doc.add_heading('Site Locations & Boundaries', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph(
        "The following maps show the geographic location and site boundaries for each project."
    )
    
    for site in site_results:
        doc.add_paragraph(f"{site.get('site_name')}:", style='Heading 2')
        
        # Site details
        doc.add_paragraph(f"Location: {site.get('location', 'N/A')}")
        doc.add_paragraph(f"Land Area: {site.get('land_acres', 0)} acres")
        doc.add_paragraph(f"IT Capacity: {site.get('it_capacity_mw', 0):.0f} MW")
        
        # Generate map (will return HTML for now)
        geojson = site.get('geojson_data')
        map_path = create_site_map_image(site, geojson)
        
        if map_path:
            doc.add_paragraph(f"Map saved to: {map_path}", style='List Bullet')
       
        doc.add_paragraph()  # Spacing


# Keep backward compatibility
def generate_word_report(site_selection: List[str], content_options: Dict, 
                         sites_list: List[Dict]) -> bytes:
    """
    Backward-compatible wrapper that calls enhanced report generator
    """
    return generate_enhanced_word_report(site_selection, content_options, sites_list)


if __name__ == "__main__":
    print("Enhanced Report Builder ready!")
    print(f"Enhanced features available: {ENHANCED_FEATURES_AVAILABLE}")
