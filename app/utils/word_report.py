"""
Comprehensive Word Document Report Generator
Creates detailed optimization reports with all site data, equipment specs, constraints, etc.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List
import io
import os


def generate_comprehensive_word_report(
    site: Dict,
    constraints: Dict,
    scenario: Dict,
    equipment_config: Dict,
    optimization_result: Dict,
    load_profile: Dict = None
) -> bytes:
    """
    Generate a comprehensive Word document report
    
    Returns:
        bytes: Word document as bytes for download
    """
    
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Energy Optimization Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"\n{site.get('Site_Name', 'Unknown Site')}\n")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(70, 130, 180)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    feasibility = "FEASIBLE ✓" if optimization_result.get('feasible') else "INFEASIBLE ✗"
    doc.add_paragraph(f"**Scenario:** {optimization_result.get('scenario_name')}")
    doc.add_paragraph(f"**Status:** {feasibility}")
    
    if optimization_result.get('feasible'):
        economics = optimization_result['economics']
        timeline = optimization_result['timeline']
        
        doc.add_paragraph(f"**LCOE:** ${economics['lcoe_mwh']:.2f}/MWh")
        doc.add_paragraph(f"**Total CAPEX:** ${economics['total_capex_m']:.1f} Million")
        doc.add_paragraph(f"**Deployment Timeline:** {timeline['timeline_months']} months ({timeline['timeline_years']:.1f} years)")
        doc.add_paragraph(f"**Annual Energy:** {economics['annual_generation_gwh']:.1f} GWh")
        
        # Key recommendation
        doc.add_heading('Recommendation', 2)
        doc.add_paragraph(
            f"This {scenario.get('Scenario_Name')} scenario provides a {timeline['deployment_speed'].lower()} "
            f"deployment option with an LCOE of ${economics['lcoe_mwh']:.2f}/MWh. "
            f"The configuration meets all site constraints and reliability requirements."
        )
    
    doc.add_page_break()
    
    # Site Information
    doc.add_heading('1. Site Information', 1)
    
    doc.add_heading('1.1 Site Overview', 2)
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'
    
    site_data = [
        ('Site Name', site.get('Site_Name', 'N/A')),
        ('Location', f"{site.get('State', 'N/A')}, {site.get('ISO', 'N/A')}"),
        ('Coordinates', f"{site.get('Latitude', 0):.4f}°, {site.get('Longitude', 0):.4f}°"),
        ('IT Capacity', f"{site.get('IT_Capacity_MW', 0)} MW"),
        ('Design PUE', f"{site.get('Design_PUE', 0):.2f}"),
        ('Total Facility Load', f"{site.get('Total_Facility_MW', 0)} MW"),
        ('Load Factor', f"{site.get('Load_Factor_Pct', 0)}%"),
        ('Altitude', f"{site.get('Altitude_ft', 0):,} ft"),
        ('Average Temperature', f"{site.get('Avg_Temp_F', 0)}°F")
    ]
    
    for i, (label, value) in enumerate(site_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    # Hard Constraints
    doc.add_heading('1.2 Hard Constraints', 2)
    
    doc.add_heading('Air Quality Permits', 3)
    air_table = doc.add_table(rows=4, cols=2)
    air_table.style = 'Light Grid Accent 1'
    
    air_constraints = [
        ('Permit Type', constraints.get('Air_Permit_Type', 'N/A')),
        ('NOx Limit', f"{constraints.get('NOx_Limit_tpy', 0)} tons/year"),
        ('CO Limit', f"{constraints.get('CO_Limit_tpy', 0)} tons/year"),
        ('Nonattainment Area', constraints.get('Nonattainment_Area', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(air_constraints):
        air_table.rows[i].cells[0].text = label
        air_table.rows[i].cells[1].text = str(value)
    
    doc.add_heading('Infrastructure Constraints', 3)
    infra_table = doc.add_table(rows=5, cols=2)
    infra_table.style = 'Light Grid Accent 1'
    
    infra_constraints = [
        ('Natural Gas Supply', f"{constraints.get('Gas_Supply_MCF_day', 0):,} MCF/day"),
        ('Gas Pipeline', constraints.get('Gas_Pipeline', 'N/A')),
        ('Grid Available', f"{constraints.get('Grid_Available_MW', 0)} MW"),
        ('Interconnection Queue', f"Position #{constraints.get('Queue_Position', 0)}"),
        ('Grid Timeline', f"{constraints.get('Estimated_Interconnection_Months', 0)} months")
    ]
    
    for i, (label, value) in enumerate(infra_constraints):
        infra_table.rows[i].cells[0].text = label
        infra_table.rows[i].cells[1].text = str(value)
    
    doc.add_heading('Land & Physical', 3)
    land_table = doc.add_table(rows=3, cols=2)
    land_table.style = 'Light Grid Accent 1'
    
    land_constraints = [
        ('Total Land', f"{constraints.get('Total_Land_Acres', 0)} acres"),
        ('Available Land', f"{constraints.get('Available_Land_Acres', 0)} acres"),
        ('Solar Feasibility', constraints.get('Solar_Feasible', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(land_constraints):
        land_table.rows[i].cells[0].text = label
        land_table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # Load Profile
    doc.add_heading('2. Load Profile & Requirements', 1)
    
    if load_profile:
        doc.add_heading('2.1 IT Workload Composition', 2)
        
        workload_mix = load_profile.get('workload_mix', {})
        if workload_mix:
            mix_table = doc.add_table(rows=len(workload_mix) + 1, cols=2)
            mix_table.style = 'Light Grid Accent 1'
            mix_table.rows[0].cells[0].text = 'Workload Type'
            mix_table.rows[0].cells[1].text = 'Percentage'
            
            for i, (workload, pct) in enumerate(workload_mix.items(), 1):
                mix_table.rows[i].cells[0].text = workload
                mix_table.rows[i].cells[1].text = f"{pct}%"
        
        doc.add_paragraph(f"**IT Capacity:** {load_profile.get('it_capacity_mw', 0)} MW")
        doc.add_paragraph(f"**PUE:** {load_profile.get('pue', 1.25):.2f}")
        doc.add_paragraph(f"**Load Factor:** {load_profile.get('load_factor', 75)}%")
    
    doc.add_heading('2.2 Reliability Requirements', 2)
    rel_table = doc.add_table(rows=3, cols=2)
    rel_table.style = 'Light Grid Accent 1'
    
    rel_requirements = [
        ('N-1 Required', 'Yes' if constraints.get('N_Minus_1_Required') == 'Yes' else 'No'),
        ('Max Transient', f"{constraints.get('Max_Transient_pct', 0)}%"),
        ('Min Spinning Reserve', f"{constraints.get('Min_Spinning_Reserve_MW', 0)} MW")
    ]
    
    for i, (label, value) in enumerate(rel_requirements):
        rel_table.rows[i].cells[0].text = label
        rel_table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # Scenario Description
    doc.add_heading('3. Scenario Configuration', 1)
    
    doc.add_paragraph(f"**Scenario:** {scenario.get('Scenario_Name', 'N/A')}")
    doc.add_paragraph(f"**Description:** {scenario.get('Description', 'N/A')}")
    doc.add_paragraph(f"**Deployment Strategy:** {scenario.get('Deployment_Strategy', 'N/A')}")
    
    doc.add_heading('3.1 Enabled Technologies', 2)
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    technologies = [
        ('Reciprocating Engines', scenario.get('Recip_Engines', 'False')),
        ('Gas Turbines', scenario.get('Gas_Turbines', 'False')),
        ('BESS', scenario.get('BESS', 'False')),
        ('Solar PV', scenario.get('Solar_PV', 'False')),
        ('Grid Connection', scenario.get('Grid_Connection', 'False'))
    ]
    
    for i, (tech, enabled) in enumerate(technologies):
        tech_table.rows[i].cells[0].text = tech
        tech_table.rows[i].cells[1].text = '✓ Enabled' if enabled == 'True' else '✗ Disabled'
    
    doc.add_page_break()
    
    # Equipment Configuration
    doc.add_heading('4. Equipment Configuration & Specifications', 1)
    
    # Reciprocating Engines
    if equipment_config.get('recip_engines'):
        doc.add_heading('4.1 Reciprocating Engines', 2)
        engines = equipment_config['recip_engines']
        
        doc.add_paragraph(f"**Quantity:** {len(engines)} units")
        
        if engines:
            engine = engines[0]
            eng_table = doc.add_table(rows=7, cols=2)
            eng_table.style = 'Light Grid Accent 1'
            
            eng_specs = [
                ('Unit Capacity', f"{engine.get('capacity_mw', 0):.1f} MW"),
                ('Capacity Factor', f"{engine.get('capacity_factor', 0):.1%}"),
                ('Heat Rate', f"{engine.get('heat_rate_btu_kwh', 0):,} Btu/kWh"),
                ('NOx Rate', f"{engine.get('nox_lb_mmbtu', 0):.3f} lb/MMBtu"),
                ('CO Rate', f"{engine.get('co_lb_mmbtu', 0):.3f} lb/MMBtu"),
                ('CAPEX', f"${engine.get('capex_per_kw', 0):,}/kW"),
                ('Total Capacity', f"{len(engines) * engine.get('capacity_mw', 0):.1f} MW")
            ]
            
            for i, (label, value) in enumerate(eng_specs):
                eng_table.rows[i].cells[0].text = label
                eng_table.rows[i].cells[1].text = str(value)
    
    # Gas Turbines
    if equipment_config.get('gas_turbines'):
        doc.add_heading('4.2 Gas Turbines', 2)
        turbines = equipment_config['gas_turbines']
        
        doc.add_paragraph(f"**Quantity:** {len(turbines)} units")
        
        if turbines:
            turbine = turbines[0]
            turb_table = doc.add_table(rows=7, cols=2)
            turb_table.style = 'Light Grid Accent 1'
            
            turb_specs = [
                ('Unit Capacity', f"{turbine.get('capacity_mw', 0):.1f} MW"),
                ('Capacity Factor', f"{turbine.get('capacity_factor', 0):.1%}"),
                ('Heat Rate', f"{turbine.get('heat_rate_btu_kwh', 0):,} Btu/kWh"),
                ('NOx Rate', f"{turbine.get('nox_lb_mmbtu', 0):.3f} lb/MMBtu"),
                ('CO Rate', f"{turbine.get('co_lb_mmbtu', 0):.3f} lb/MMBtu"),
                ('CAPEX', f"${turbine.get('capex_per_kw', 0):,}/kW"),
                ('Total Capacity', f"{len(turbines) * turbine.get('capacity_mw', 0):.1f} MW")
            ]
            
            for i, (label, value) in enumerate(turb_specs):
                turb_table.rows[i].cells[0].text = label
                turb_table.rows[i].cells[1].text = str(value)
    
    # BESS
    if equipment_config.get('bess'):
        doc.add_heading('4.3 Battery Energy Storage System (BESS)', 2)
        bess_units = equipment_config['bess']
        
        doc.add_paragraph(f"**Quantity:** {len(bess_units)} units")
        
        if bess_units:
            bess = bess_units[0]
            bess_table = doc.add_table(rows=5, cols=2)
            bess_table.style = 'Light Grid Accent 1'
            
            total_energy = len(bess_units) * bess.get('energy_mwh', 0)
            total_power = len(bess_units) * bess.get('power_mw', 0)
            
            bess_specs = [
                ('Unit Energy', f"{bess.get('energy_mwh', 0):.1f} MWh"),
                ('Unit Power', f"{bess.get('power_mw', 0):.1f} MW"),
                ('Total Energy', f"{total_energy:.1f} MWh"),
                ('Total Power', f"{total_power:.1f} MW"),
                ('CAPEX', f"${bess.get('capex_per_kwh', 0):,}/kWh")
            ]
            
            for i, (label, value) in enumerate(bess_specs):
                bess_table.rows[i].cells[0].text = label
                bess_table.rows[i].cells[1].text = str(value)
    
    # Solar
    if equipment_config.get('solar_mw_dc'):
        doc.add_heading('4.4 Solar PV', 2)
        solar_table = doc.add_table(rows=3, cols=2)
        solar_table.style = 'Light Grid Accent 1'
        
        solar_specs = [
            ('DC Capacity', f"{equipment_config.get('solar_mw_dc', 0):.1f} MW"),
            ('Land Required', f"{equipment_config.get('solar_mw_dc', 0) * 4.25:.1f} acres"),
            ('CAPEX', f"${equipment_config.get('solar_capex_per_w', 0):.2f}/W")
        ]
        
        for i, (label, value) in enumerate(solar_specs):
            solar_table.rows[i].cells[0].text = label
            solar_table.rows[i].cells[1].text = str(value)
    
    # Grid
    if equipment_config.get('grid_import_mw'):
        doc.add_heading('4.5 Grid Connection', 2)
        doc.add_paragraph(f"**Import Capacity:** {equipment_config.get('grid_import_mw', 0):.1f} MW")
    
    doc.add_page_break()
    
    # 8760 Dispatch & Power Quality Visualization
    if optimization_result.get('feasible'):
        doc.add_heading('5. 8760 Dispatch & Power Quality Analysis', 1)
        
        # Run dispatch simulation to get actual 8760 data
        dispatch_data = None
        try:
            from app.utils.dispatch_simulation import dispatch_equipment, generate_8760_load_profile
            
            # Generate load profile
            load_mw = site.get('Total_Facility_MW', 200)
            load_factor = site.get('Load_Factor_Pct', 70) / 100
            load_profile_array = generate_8760_load_profile(load_mw, load_factor)
            
            # Run dispatch
            dispatch_results = dispatch_equipment(load_profile_array, equipment_config, bess_available=True)
            dispatch_data = dispatch_results  # This has all the hourly arrays
        except Exception as e:
            # If dispatch fails, charts will fall back to synthetic data
            pass
        
        try:
            from app.utils.report_charts import (
                create_8760_dispatch_chart,
                create_emissions_chart,
                create_bess_soc_chart,
                create_deployment_timeline_chart
            )
            
            # Generate and embed dispatch chart
            doc.add_heading('5.1 Hourly Dispatch Visualization', 2)
            doc.add_paragraph("The following chart shows the equipment dispatch stack for the first week (168 hours).")
            
            dispatch_chart = create_8760_dispatch_chart(equipment_config, site, dispatch_data=dispatch_data)
            if dispatch_chart and os.path.exists(dispatch_chart):
                doc.add_picture(dispatch_chart, width=Inches(6.5))
                os.remove(dispatch_chart)  # Clean up temp file
            
            # BESS State of Charge
            if equipment_config.get('bess'):
                doc.add_heading('5.2 BESS State of Charge', 2)
                doc.add_paragraph("Battery state of charge over the first week, showing daily charge/discharge cycles.")
                
                bess_chart = create_bess_soc_chart(equipment_config)
                if bess_chart and os.path.exists(bess_chart):
                    doc.add_picture(bess_chart, width=Inches(6))
                    os.remove(bess_chart)
            
            # Emissions Analysis
            doc.add_heading('5.3 Hourly Emissions Analysis', 2)
            doc.add_paragraph("NOx and CO emissions from generators, compared against annual average limits.")
            
            emissions_chart = create_emissions_chart(equipment_config, constraints)
            if emissions_chart and os.path.exists(emissions_chart):
                doc.add_picture(emissions_chart, width=Inches(6.5))
                os.remove(emissions_chart)
            
            # Deployment Timeline
            timeline = optimization_result.get('timeline', {})
            doc.add_heading('5.4 Equipment Deployment Timeline', 2)
            doc.add_paragraph("Gantt chart showing deployment phases and critical path.")
            
            timeline_chart = create_deployment_timeline_chart(timeline)
            if timeline_chart and os.path.exists(timeline_chart):
                doc.add_picture(timeline_chart, width=Inches(6))
                os.remove(timeline_chart)
                
        except Exception as e:
            doc.add_paragraph(f"Note: Visualization generation encountered an error: {str(e)}")
    
    doc.add_page_break()
    
    # Optimization Results
    if optimization_result.get('feasible'):
        doc.add_heading('6. Optimization Results', 1)
        
        economics = optimization_result['economics']
        timeline = optimization_result['timeline']
        metrics = optimization_result['metrics']
        
        doc.add_heading('6.1 Economic Analysis', 2)
        econ_table = doc.add_table(rows=6, cols=2)
        econ_table.style = 'Light Grid Accent 1'
        
        econ_data = [
            ('LCOE', f"${economics['lcoe_mwh']:.2f}/MWh"),
            ('Total CAPEX', f"${economics['total_capex_m']:.2f} Million"),
            ('Annual O&M', f"${economics['annual_opex_m']:.2f} Million"),
            ('Annual Fuel Cost', f"${economics['annual_fuel_cost_m']:.2f} Million"),
            ('Annual Generation', f"{economics['annual_generation_gwh']:.1f} GWh"),
            ('Capacity Factor', f"{economics['capacity_factor_pct']:.1f}%")
        ]
        
        for i, (label, value) in enumerate(econ_data):
            econ_table.rows[i].cells[0].text = label
            econ_table.rows[i].cells[1].text = str(value)
        
        doc.add_heading('6.2 Deployment Timeline', 2)
        time_table = doc.add_table(rows=3, cols=2)
        time_table.style = 'Light Grid Accent 1'
        
        time_data = [
            ('Total Timeline', f"{timeline['timeline_months']} months ({timeline['timeline_years']:.1f} years)"),
            ('Deployment Speed', timeline['deployment_speed']),
            ('Critical Path', timeline['critical_path'])
        ]
        
        for i, (label, value) in enumerate(time_data):
            time_table.rows[i].cells[0].text = label
            time_table.rows[i].cells[1].text = str(value)
        
        doc.add_heading('6.3 Performance Metrics', 2)
        metrics_table = doc.add_table(rows=3, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        # Calculate capacities from equipment_config
        total_cap = 0
        btm_cap = 0
        grid_cap = equipment_config.get('grid_import_mw', 0)
        
        # Reciprocating engines
        if equipment_config.get('recip_engines'):
            recip_total = sum(e.get('capacity_mw', 0) for e in equipment_config['recip_engines'])
            total_cap += recip_total
            btm_cap += recip_total
        
        # Gas turbines
        if equipment_config.get('gas_turbines'):
            turbine_total = sum(t.get('capacity_mw', 0) for t in equipment_config['gas_turbines'])
            total_cap += turbine_total
            btm_cap += turbine_total
        
        # BESS (power capacity)
        if equipment_config.get('bess'):
            bess_total = sum(b.get('power_mw', 0) for b in equipment_config['bess'])
            total_cap += bess_total
            btm_cap += bess_total
        
        # Solar (DC capacity)
        if equipment_config.get('solar_mw_dc'):
            solar_dc = equipment_config.get('solar_mw_dc', 0)
            total_cap += solar_dc
            btm_cap += solar_dc
        
        # Grid
        total_cap += grid_cap
        
        metrics_data = [
            ('Total Capacity', f"{total_cap:.1f} MW"),
            ('BTM Capacity', f"{btm_cap:.1f} MW"),
            ('Grid Capacity', f"{grid_cap:.1f} MW")
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = str(value)
        
        # RAM Analysis (if available)
        if optimization_result.get('ram_analysis'):
            doc.add_heading('6.4 Reliability, Availability & Maintainability (RAM)', 2)
            
            ram = optimization_result['ram_analysis']
            
            if ram.get('equipment'):
                doc.add_paragraph("**Equipment Reliability Metrics:**")
                
                ram_table = doc.add_table(rows=len(ram['equipment']) + 1, cols=5)
                ram_table.style = 'Light Grid Accent 1'
                
                # Header
                ram_table.rows[0].cells[0].text = 'Equipment Type'
                ram_table.rows[0].cells[1].text = 'Quantity'
                ram_table.rows[0].cells[2].text = 'System Availability'
                ram_table.rows[0].cells[3].text = 'MTBF (hrs)'
                ram_table.rows[0].cells[4].text = 'MTTR (hrs)'
                
                for i, eq in enumerate(ram['equipment'], 1):
                    ram_table.rows[i].cells[0].text = eq.get('Type', '')
                    ram_table.rows[i].cells[1].text = str(eq.get('Count', 0))
                    ram_table.rows[i].cells[2].text = str(eq.get('System Availability', 'N/A'))
                    ram_table.rows[i].cells[3].text = str(eq.get('MTBF (hrs)', 'N/A'))
                    ram_table.rows[i].cells[4].text = str(eq.get('MTTR (hrs)', 'N/A'))
            
            # System-level metrics
            if ram.get('system_availability'):
                doc.add_paragraph(f"\n**System Availability:** {ram.get('system_availability', 0):.4f} ({ram.get('system_availability', 0)*100:.2f}%)")
            if ram.get('expected_outages_per_year'):
                doc.add_paragraph(f"**Expected Outages per Year:** {ram.get('expected_outages_per_year', 0):.1f}")
            if ram.get('expected_downtime_hours_per_year'):
                doc.add_paragraph(f"**Expected Downtime:** {ram.get('expected_downtime_hours_per_year', 0):.1f} hours/year")
        
        # Transient Analysis (if available)
        if optimization_result.get('transient_analysis'):
            doc.add_heading('6.5 Transient & Power Quality Analysis', 2)
            
            transient = optimization_result['transient_analysis']
            pq_metrics = transient.get('pq_metrics', {})
            
            if pq_metrics:
                doc.add_paragraph("**Power Quality Metrics (20% Step Change Event):**")
                
                pq_table = doc.add_table(rows=8, cols=2)
                pq_table.style = 'Light Grid Accent 1'
                
                pq_data = [
                    ('Max Frequency Deviation', f"{pq_metrics.get('max_frequency_deviation_hz', 0):.3f} Hz"),
                    ('Frequency Nadir', f"{pq_metrics.get('frequency_nadir_hz', 60.0):.2f} Hz"),
                    ('Frequency Zenith', f"{pq_metrics.get('frequency_zenith_hz', 60.0):.2f} Hz"),
                    ('Max Ramp Rate', f"{pq_metrics.get('max_ramp_rate_mw_s', 0):.2f} MW/s"),
                    ('Avg Ramp Rate', f"{pq_metrics.get('avg_ramp_rate_mw_s', 0):.2f} MW/s"),
                    ('BESS Max Response', f"{pq_metrics.get('bess_max_response_mw', 0):.1f} MW"),
                    ('Time to Stabilize', f"{pq_metrics.get('time_to_stabilize_s', 0):.0f} seconds"),
                    ('Transient Severity', pq_metrics.get('transient_severity', 'Unknown'))
                ]
                
                for i, (label, value) in enumerate(pq_data):
                    pq_table.rows[i].cells[0].text = label
                    pq_table.rows[i].cells[1].text = str(value)
                
                doc.add_paragraph("\n**Analysis:** This transient simulation models a 20% sudden load change to evaluate system response and power quality impacts.")
                
                # Add transient visualization charts if data available
                transient_data = transient.get('transient_data', {})
                if transient_data:
                    try:
                        from app.utils.transient_charts import (
                            create_transient_response_chart,
                            create_load_rate_of_change_chart,
                            create_frequency_deviation_chart,
                            create_workload_step_change_chart
                        )
                        
                        doc.add_paragraph("\n**Transient Visualizations:**")
                        
                        # Workload step change
                        step_chart = create_workload_step_change_chart(transient_data)
                        if step_chart and os.path.exists(step_chart):
                            doc.add_paragraph("*Workload Step Change Event:*")
                            doc.add_picture(step_chart, width=Inches(6))
                            os.remove(step_chart)
                        
                        # Transient response (Load/Gen/BESS)
                        response_chart = create_transient_response_chart(transient_data)
                        if response_chart and os.path.exists(response_chart):
                            doc.add_paragraph("*System Response (Load, Generator, BESS):*")
                            doc.add_picture(response_chart, width=Inches(6))
                            os.remove(response_chart)
                        
                        # Load rate of change
                        rate_chart = create_load_rate_of_change_chart(transient_data)
                        if rate_chart and os.path.exists(rate_chart):
                            doc.add_paragraph("*Load Rate of Change (dP/dt):*")
                            doc.add_picture(rate_chart, width=Inches(6))
                            os.remove(rate_chart)
                        
                        # Frequency deviation
                        freq_chart = create_frequency_deviation_chart(transient_data)
                        if freq_chart and os.path.exists(freq_chart):
                            doc.add_paragraph("*Frequency Deviation:*")
                            doc.add_picture(freq_chart, width=Inches(6))
                            os.remove(freq_chart)
                            
                    except Exception as e:
                        doc.add_paragraph(f"*Note: Some transient charts unavailable: {str(e)}*")
    
    # Constraint Compliance
    doc.add_page_break()
    doc.add_heading('7. Constraint Compliance', 1)
    
    violations = optimization_result.get('violations', [])
    warnings = optimization_result.get('warnings', [])
    
    if not violations:
        doc.add_paragraph("✓ All hard constraints satisfied")
    else:
        doc.add_heading('Violations', 2)
        for v in violations:
            doc.add_paragraph(f"• {v}", style='List Bullet')
    
    if warnings:
        doc.add_heading('Warnings', 2)
        for w in warnings:
            doc.add_paragraph(f"• {w}", style='List Bullet')
    
    # Methodology
    doc.add_page_break()
    doc.add_heading('8. Methodology', 1)
    
    doc.add_paragraph(
        "This analysis uses a multi-criteria optimization approach to determine the optimal energy "
        "configuration for the datacenter site. The methodology includes:"
    )
    
    methodology_steps = [
        "Load modeling based on IT workload characteristics and facility PUE",
        "Equipment sizing using heuristic algorithms and site-specific constraints",
        "Constraint validation for air quality permits, gas supply, land use, and reliability",
        "Economic analysis using NPV method with 20-year project life",
        "Deployment timeline analysis considering equipment lead times and permitting",
        "Multi-objective optimization balancing LCOE, deployment speed, and reliability"
    ]
    
    for step in methodology_steps:
        doc.add_paragraph(step, style='List Number')
    
    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix: Reference Data', 1)
    
    # A. Equipment Database Reference
    doc.add_heading('A. Equipment Specifications Reference', 2)
    
    doc.add_heading('Reciprocating Engines', 3)
    if equipment_config.get('recip_engines') and len(equipment_config['recip_engines']) > 0:
        recip = equipment_config['recip_engines'][0]
        recip_ref_table = doc.add_table(rows=7, cols=2)
        recip_ref_table.style = 'Light Grid Accent 1'
        
        recip_ref_data = [
            ('Unit Capacity', f"{recip.get('capacity_mw', 0):.1f} MW"),
            ('Heat Rate', f"{recip.get('heat_rate_btu_kwh', 0):,} Btu/kWh"),
            ('NOx Emission Rate', f"{recip.get('nox_lb_mmbtu', 0):.4f} lb/MMBtu"),
            ('CO Emission Rate', f"{recip.get('co_lb_mmbtu', 0):.4f} lb/MMBtu"),
            ('CAPEX', f"${recip.get('capex_per_kw', 0):,}/kW"),
            ('Lead Time', "18-24 months (typical)"),
            ('Operational Life', "20-25 years")
        ]
        
        for i, (label, value) in enumerate(recip_ref_data):
            recip_ref_table.rows[i].cells[0].text = label
            recip_ref_table.rows[i].cells[1].text = str(value)
    
    doc.add_heading('Gas Turbines', 3)
    if equipment_config.get('gas_turbines') and len(equipment_config['gas_turbines']) > 0:
        turbine = equipment_config['gas_turbines'][0]
        turbine_ref_table = doc.add_table(rows=7, cols=2)
        turbine_ref_table.style = 'Light Grid Accent 1'
        
        turbine_ref_data = [
            ('Unit Capacity', f"{turbine.get('capacity_mw', 0):.1f} MW"),
            ('Heat Rate', f"{turbine.get('heat_rate_btu_kwh', 0):,} Btu/kWh"),
            ('NOx Emission Rate', f"{turbine.get('nox_lb_mmbtu', 0):.4f} lb/MMBtu"),
            ('CO Emission Rate', f"{turbine.get('co_lb_mmbtu', 0):.4f} lb/MMBtu"),
            ('CAPEX', f"${turbine.get('capex_per_kw', 0):,}/kW"),
            ('Lead Time', "24-36 months (typical)"),
            ('Operational Life', "25-30 years")
        ]
        
        for i, (label, value) in enumerate(turbine_ref_data):
            turbine_ref_table.rows[i].cells[0].text = label
            turbine_ref_table.rows[i].cells[1].text = str(value)
    
    # B. Site Constraints Summary
    doc.add_heading('B. Site Constraints Summary', 2)
    
    constraint_summary_table = doc.add_table(rows=10, cols=3)
    constraint_summary_table.style = 'Light Grid Accent 1'
    
    # Header
    constraint_summary_table.rows[0].cells[0].text = 'Constraint Type'
    constraint_summary_table.rows[0].cells[1].text = 'Limit'
    constraint_summary_table.rows[0].cells[2].text = 'Status'
    
    constraint_summary_data = [
        ('NOx Emissions', f"{constraints.get('NOx_Limit_tpy', 0)} tpy", 'Hard Limit'),
        ('CO Emissions', f"{constraints.get('CO_Limit_tpy', 0)} tpy", 'Hard Limit'),
        ('Natural Gas Supply', f"{constraints.get('Gas_Supply_MCF_day', 0):,} MCF/day", 'Hard Limit'),
        ('Grid Capacity', f"{constraints.get ('Grid_Available_MW', 0)} MW", 'Hard Limit'),
        ('Available Land', f"{constraints.get('Available_Land_Acres', 0)} acres", 'Hard Limit'),
        ('N-1 Reliability', constraints.get('N_Minus_1_Required', 'No'), 'Hard Requirement'),
        ('Max Transient', f"{constraints.get('Max_Transient_pct', 0)}%", 'Soft Guideline'),
        ('Interconnection Timeline', f"{constraints.get('Estimated_Interconnection_Months', 0)} months", 'Schedule'),
        ('Permit Type', constraints.get('Air_Permit_Type', 'N/A'), 'Regulatory')
    ]
    
    for i, (ctype, limit, status) in enumerate(constraint_summary_data, 1):
        constraint_summary_table.rows[i].cells[0].text = ctype
        constraint_summary_table.rows[i].cells[1].text = str(limit)
        constraint_summary_table.rows[i].cells[2].text = status
    
    # C. Optimization Parameters
    doc.add_heading('C. Optimization Parameters', 2)
    
    doc.add_paragraph("**Optimization Method:** Sequential Least Squares Programming (SLSQP)")
    doc.add_paragraph("**Solver:** SciPy optimize.minimize")
    doc.add_paragraph("**Decision Variables:** Equipment quantities, capacity factors, renewable sizing")
    doc.add_paragraph("**Objective Function:** Weighted combination of LCOE, deployment timeline, and emissions")
    doc.add_paragraph("**Constraint Handling:** Penalty-based approach with hard limit enforcement")
    
    # D. Glossary
    doc.add_heading('D. Glossary of Terms', 2)
    
    glossary_table = doc.add_table(rows=16, cols=2)
    glossary_table.style = 'Light Grid Accent 1'
    
    glossary_data = [
        ('LCOE', 'Levelized Cost of Energy ($/MWh) - Total lifetime cost divided by total energy produced'),
        ('CAPEX', 'Capital Expenditure - Upfront equipment and installation costs'),
        ('OPEX', 'Operating Expenditure - Annual maintenance, labor, and operating costs'),
        ('PUE', 'Power Usage Effectiveness - Ratio of total facility power to IT equipment power'),
        ('BESS', 'Battery Energy Storage System'),
        ('BTM', 'Behind-the-Meter - On-site generation not connected to grid'),
        ('IFOM', 'In-Front-of-Meter - Grid-connected generation'),
        ('N-1', 'Reliability criteria - System must function with any single equipment failure'),
        ('MCF', 'Thousand Cubic Feet (natural gas volume measurement)'),
        ('tpy', 'Tons per year (emissions measurement)'),
        ('NOx', 'Nitrogen Oxides - Regulated air pollutant'),
        ('CO', 'Carbon Monoxide - Regulated air pollutant'),
        ('SOC', 'State of Charge - Battery energy level as percentage of capacity'),
        ('MW', 'Megawatt - Unit of power (1,000 kW)'),
        ('MWh', 'Megawatt-hour - Unit of energy')
    ]
    
    for i, (term, definition) in enumerate(glossary_data):
        glossary_table.rows[i].cells[0].text = term
        glossary_table.rows[i].cells[1].text = definition
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()
