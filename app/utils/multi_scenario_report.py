"""
Multi-Scenario Report Generator
Creates comprehensive Word reports with all scenarios
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List
import io


def generate_multi_scenario_report(
    site: Dict,
    constraints: Dict,
    all_results: List[Dict],
    load_profile: Dict = None
) -> bytes:
    """
    Generate a comprehensive Word document with ALL scenarios
    
    Args:
        site: Site information
        constraints: Site constraints
        all_results: List of all scenario results (feasible + infeasible)
        load_profile: Load profile data
    
    Returns:
        bytes: Word document as bytes
    """
    
    doc = Document()
    
    # ========== TITLE PAGE ==========
    title = doc.add_heading('Multi-Scenario Optimization Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"\\n{site.get('Site_Name', 'Unknown Site')}\\n")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(70, 130, 180)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading('Executive Summary', 1)
    
    feasible_scenarios = [r for r in all_results if r.get('feasible')]
    infeasible_scenarios = [r for r in all_results if not r.get('feasible')]
    
    doc.add_paragraph(f"**Total Scenarios Evaluated:** {len(all_results)}")
    doc.add_paragraph(f"**Feasible Scenarios:** {len(feasible_scenarios)}")
    doc.add_paragraph(f"**Infeasible Scenarios:** {len(infeasible_scenarios)}")
    
    if feasible_scenarios:
        best_scenario = feasible_scenarios[0]  # Already ranked
        doc.add_paragraph(f"**Recommended Scenario:** {best_scenario.get('scenario_name')}")
        doc.add_paragraph(f"**LCOE:** ${best_scenario['economics']['lcoe_mwh']:.2f}/MWh")
        doc.add_paragraph(f"**Total CAPEX:** ${best_scenario['economics']['total_capex_m']:.1f} Million")
        doc.add_paragraph(f"**Deployment:** {best_scenario['timeline']['timeline_months']} months")
    
    # ========== SCENARIO COMPARISON TABLE ==========
    doc.add_heading('Scenario Comparison', 1)
    
    # Create comparison table
    num_rows = len(all_results) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Rank', 'Scenario', 'Feasible', 'LCOE ($/MWh)', 'CAPEX ($M)', 'Timeline (mo)', 'Constraint Violations']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for i, result in enumerate(all_results, 1):
        row = table.rows[i]
        row.cells[0].text = str(result.get('rank', 999))
        row.cells[1].text = result.get('scenario_name', 'Unknown')
        row.cells[2].text = '✓' if result.get('feasible') else '✗'
        
        if result.get('feasible'):
            row.cells[3].text = f"${result['economics']['lcoe_mwh']:.2f}"
            row.cells[4].text = f"${result['economics']['total_capex_m']:.1f}"
            row.cells[5].text = str(result['timeline']['timeline_months'])
            row.cells[6].text = '-'
        else:
            row.cells[3].text = 'N/A'
            row.cells[4].text = 'N/A'
            row.cells[5].text = 'N/A'
            violations = result.get('violations', [])
            if violations:
                row.cells[6].text = '; '.join(violations[:2])
                if len(violations) > 2:
                    row.cells[6].text += f" (+{len(violations)-2} more)"
            else:
                row.cells[6].text = 'Unknown'
    
    doc.add_page_break()
    
    # ========== SITE INFORMATION (ONCE) ==========
    doc.add_heading('Site Information', 1)
    
    doc.add_heading('Site Overview', 2)
    site_table = doc.add_table(rows=7, cols=2)
    site_table.style = 'Light Grid Accent 1'
    
    site_data = [
        ('Site Name', site.get('Site_Name', 'N/A')),
        ('Location', f"{site.get('State', 'N/A')}, {site.get('ISO', 'N/A')}"),
        ('IT Capacity', f"{site.get('IT_Capacity_MW', 0)} MW"),
        ('Design PUE', f"{site.get('Design_PUE', 0):.2f}"),
        ('Total Facility Load', f"{site.get('Total_Facility_MW', 0)} MW"),
        ('Load Factor', f"{site.get('Load_Factor_Pct', 0)}%"),
        ('Coordinates', f"{site.get('Latitude', 0):.4f}°, {site.get('Longitude', 0):.4f}°")
    ]
    
    for i, (label, value) in enumerate(site_data):
        site_table.rows[i].cells[0].text = label
        site_table.rows[i].cells[1].text = str(value)
    
    # Site Constraints
    doc.add_heading('Site Constraints', 2)
    
    const_table = doc.add_table(rows=9, cols=2)
    const_table.style = 'Light Grid Accent 1'
    
    constraint_data = [
        ('NOx Limit', f"{constraints.get('NOx_Limit_tpy', 0)} tons/year"),
        ('CO Limit', f"{constraints.get('CO_Limit_tpy', 0)} tons/year"),
        ('Gas Supply', f"{constraints.get('Gas_Supply_MCF_day', 0):,} MCF/day"),
        ('Grid Available', f"{constraints.get('Grid_Available_MW', 0)} MW"),
        ('Grid Timeline', f"{constraints.get('Estimated_Interconnection_Months', 0)} months"),
        ('Available Land', f"{constraints.get('Available_Land_Acres', 0)} acres"),
        ('N-1 Required', constraints.get('N_Minus_1_Required', 'No')),
        ('Max Transient', f"{constraints.get('Max_Transient_pct', 0)}%"),
        ('Permit Type', constraints.get('Air_Permit_Type', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(constraint_data):
        const_table.rows[i].cells[0].text = label
        const_table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # ========== DETAILED SCENARIO SECTIONS ==========
    for idx, result in enumerate(all_results, 1):
        doc.add_heading(f"{idx}. {result.get('scenario_name', 'Unknown Scenario')}", 1)
        
        # Feasibility Status
        if result.get('feasible'):
            status_para = doc.add_paragraph()
            status_run = status_para.add_run("✓ FEASIBLE")
            status_run.font.bold = True
            status_run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            status_para = doc.add_paragraph()
            status_run = status_para.add_run("✗ INFEASIBLE")
            status_run.font.bold = True
            status_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Show constraint violations
            violations = result.get('violations', [])
            if violations:
                doc.add_heading('Constraint Violations', 2)
                for v in violations:
                    doc.add_paragraph(f"• {v}", style='List Bullet')
            
            doc.add_page_break()
            continue
        
        # For feasible scenarios, show full details
        economics = result['economics']
        timeline = result['timeline']
        equipment_config = result.get('equipment_config', {})
        
        # Economic Analysis
        doc.add_heading('Economic Analysis', 2)
        econ_table = doc.add_table(rows=6, cols=2)
        econ_table.style = 'Light Grid Accent 1'
        
        econ_data = [
            ('LCOE', f"${economics['lcoe_mwh']:.2f}/MWh"),
            ('Total CAPEX', f"${economics['total_capex_m']:.2f} Million"),
            ('Annual O&M', f"${economics['annual_opex_m']:.2f} Million"),
            ('Annual Fuel Cost', f"${economics['annual_fuel_cost_m']:.2f} Million"),
            ('Annual Generation', f"{economics['annual_generation_gwh']:.1f} GWh"),
            ('Capacity Factor', f"{economics['capacity_factor_pct']:.1f}%")
        ]
        
        for i, (label, value) in enumerate(econ_data):
            econ_table.rows[i].cells[0].text = label
            econ_table.rows[i].cells[1].text = str(value)
        
        # Deployment Timeline
        doc.add_heading('Deployment Timeline', 2)
        time_table = doc.add_table(rows=3, cols=2)
        time_table.style = 'Light Grid Accent 1'
        
        time_data = [
            ('Total Timeline', f"{timeline['timeline_months']} months ({timeline['timeline_years']:.1f} years)"),
            ('Deployment Speed', timeline['deployment_speed']),
            ('Critical Path', timeline['critical_path'])
        ]
        
        for i, (label, value) in enumerate(time_data):
            time_table.rows[i].cells[0].text = label
            time_table.rows[i].cells[1].text = str(value)
        
        # Equipment Configuration
        doc.add_heading('Equipment Configuration', 2)
        
        # Calculate totals
        total_cap = 0
        btm_cap = 0
        grid_cap = equipment_config.get('grid_import_mw', 0)
        
        # Recip engines
        if equipment_config.get('recip_engines'):
            recip_total = sum(e.get('capacity_mw', 0) for e in equipment_config['recip_engines'])
            total_cap += recip_total
            btm_cap += recip_total
            doc.add_paragraph(f"**Reciprocating Engines:** {len(equipment_config['recip_engines'])} units, {recip_total:.1f} MW total")
        
        # Gas turbines
        if equipment_config.get('gas_turbines'):
            turbine_total = sum(t.get('capacity_mw', 0) for t in equipment_config['gas_turbines'])
            total_cap += turbine_total
            btm_cap += turbine_total
            doc.add_paragraph(f"**Gas Turbines:** {len(equipment_config['gas_turbines'])} units, {turbine_total:.1f} MW total")
        
        # BESS
        if equipment_config.get('bess'):
            bess_total = sum(b.get('power_mw', 0) for b in equipment_config['bess'])
            total_cap += bess_total
            btm_cap += bess_total
            doc.add_paragraph(f"**BESS:** {len(equipment_config['bess'])} units, {bess_total:.1f} MW total")
        
        # Solar
        if equipment_config.get('solar_mw_dc'):
            solar_dc = equipment_config.get('solar_mw_dc', 0)
            total_cap += solar_dc
            btm_cap += solar_dc
            doc.add_paragraph(f"**Solar PV:** {solar_dc:.1f} MW DC")
        
        # Grid
        if grid_cap > 0:
            total_cap += grid_cap
            doc.add_paragraph(f"**Grid Connection:** {grid_cap:.1f} MW")
        
        doc.add_paragraph(f"\\n**Total Capacity:** {total_cap:.1f} MW")
        doc.add_paragraph(f"**BTM Capacity:** {btm_cap:.1f} MW")
        
        doc.add_page_break()
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()
