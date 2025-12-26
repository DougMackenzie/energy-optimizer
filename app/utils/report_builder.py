"""
Report Builder Module
Generate Word documents with customizable content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from typing import List, Dict
import pandas as pd

def generate_word_report(site_selection: List[str], content_options: Dict, 
                         sites_list: List[Dict]) -> bytes:
    """
    Generate Word document report
    
    Args:
        site_selection: List of selected site names
        content_options: Dict of content sections to include (True/False)
        sites_list: Full list of sites
    
    Returns:
        Word document as bytes
    """
    # Create new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Energy Optimization Report"
    doc.core_properties.author = "Antigravity Energy Optimizer"
    doc.core_properties.created = datetime.now()
    
    # =============================================================================
    # Title Page
    # =============================================================================
    if content_options.get('include_exec_overview', True):
        add_title_page(doc, site_selection)
        doc.add_page_break()
    
    # =============================================================================
    # Executive Summary
    # =============================================================================
    if content_options.get('include_exec_metrics', True):
        add_executive_summary(doc, site_selection, sites_list)
        doc.add_page_break()
    
    # =============================================================================
    # Financial Analysis
    # =============================================================================
    if content_options.get('include_cash_flow', True) or content_options.get('include_npv_irr', True):
        add_financial_analysis(doc, site_selection, sites_list, content_options)
        doc.add_page_break()
    
    # =============================================================================
    # Technical Analysis
    # =============================================================================
    if content_options.get('include_equipment', True) or content_options.get('include_optimization', True):
        add_technical_analysis(doc, site_selection, sites_list, content_options)
        doc.add_page_break()
    
    # =============================================================================
    # Site Information
    # =============================================================================
    if content_options.get('include_location_map', True):
        add_site_information(doc, site_selection, sites_list)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def add_title_page(doc: Document, site_selection: List[str]):
    """Add title page to document"""
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ENERGY OPTIMIZATION REPORT")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph()
    
    # Subtitle
    if "📊 Entire Portfolio" in site_selection:
        subtitle_text = "Portfolio Analysis"
    elif len(site_selection) > 1:
        subtitle_text = f"{len(site_selection)} Sites Analysis"
    else:
        subtitle_text = site_selection[0] if site_selection else "Site Analysis"
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(subtitle_text)
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Prepared by
    prep_para = doc.add_paragraph()
    prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_para.add_run("Prepared by: Antigravity Energy Optimizer")
    run.font.size = Pt(11)
    run.font.italic = True


def add_executive_summary(doc: Document, site_selection: List[str], sites_list: List[Dict]):
    """Add executive summary section"""
    
    # Section heading
    heading = doc.add_heading('Executive Summary', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    # Load data
    from app.utils.site_backend import load_site_stage_result
    from app.utils.financial_calculations import calculate_site_financials, calculate_portfolio_metrics
    
    portfolio_data = []
    
    for site in sites_list:
        if "📊 Entire Portfolio" not in site_selection and site.get('name') not in site_selection:
            continue
        
        site_name = site.get('name', 'Unknown')
        
        # Get latest results
        for stage in ['detailed', 'preliminary', 'concept', 'screening']:
            result = load_site_stage_result(site_name, stage)
            if result and result.get('complete'):
                financials = calculate_site_financials(site, result)
                portfolio_data.append({
                    'site': site_name,
                    'stage': stage.capitalize(),
                    'capacity_mw': site.get('it_capacity_mw', 0),
                    **financials
                })
                break
    
    if portfolio_data:
        metrics = calculate_portfolio_metrics(portfolio_data)
        
        # Overview paragraph
        doc.add_paragraph(
            f"This report analyzes {len(portfolio_data)} site(s) with a total IT capacity of "
            f"{metrics['total_capacity_mw']:.0f} MW. The portfolio demonstrates a weighted average "
            f"LCOE of ${metrics['weighted_lcoe']:.1f}/MWh with a total NPV of ${metrics['total_npv']:.1f}M."
        )
        
        doc.add_paragraph()
        
        # Key metrics table
        doc.add_paragraph("Key Portfolio Metrics:", style='Heading 2')
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        metrics_data = [
            ("Total Portfolio NPV", f"${metrics['total_npv']:.1f}M"),
            ("Weighted Average LCOE", f"${metrics['weighted_lcoe']:.1f}/MWh"),
            ("Total CapEx Required", f"${metrics['total_capex']:.1f}M"),
            ("Portfolio IRR", f"{metrics['portfolio_irr']:.1f}%"),
            ("Total IT Capacity", f"{metrics['total_capacity_mw']:.0f} MW")
        ]
        
        for i, (metric, value) in enumerate(metrics_data):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = value
            
            # Bold the metric names
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True


def add_financial_analysis(doc: Document, site_selection: List[str], 
                           sites_list: List[Dict], content_options: Dict):
    """Add financial analysis section"""
    
    # Section heading
    heading = doc.add_heading('Financial Analysis', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    from app.utils.site_backend import load_site_stage_result
    from app.utils.financial_calculations import calculate_site_financials
    
    # Site-by-site financials table
    if content_options.get('include_npv_irr', True):
        doc.add_paragraph("Site Financial Summary:", style='Heading 2')
        
        # Collect data
        financial_data = []
        for site in sites_list:
            if "📊 Entire Portfolio" not in site_selection and site.get('name') not in site_selection:
                continue
            
            site_name = site.get('name', 'Unknown')
            
            for stage in ['detailed', 'preliminary', 'concept', 'screening']:
                result = load_site_stage_result(site_name, stage)
                if result and result.get('complete'):
                    financials = calculate_site_financials(site, result)
                    financial_data.append({
                        'Site': site_name,
                        'Stage': stage.capitalize(),
                        'CapEx ($M)': f"${financials['capex_m']:.1f}",
                        'OpEx ($/yr)': f"${financials['opex_annual_m']:.1f}M",
                        'NPV ($M)': f"${financials['npv_m']:.1f}",
                        'IRR (%)': f"{financials['irr_pct']:.1f}",
                        'LCOE ($/MWh)': f"${financials['lcoe']:.1f}"
                    })
                    break
        
        if financial_data:
            df = pd.DataFrame(financial_data)
            
            # Create table
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'
            
            # Headers
            for i, col in enumerate(df.columns):
                cell = table.rows[0].cells[i]
                cell.text = col
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for row_idx, row_data in df.iterrows():
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx + 1].cells[col_idx].text = str(value)


def add_technical_analysis(doc: Document, site_selection: List[str], 
                           sites_list: List[Dict], content_options: Dict):
    """Add technical analysis section"""
    
    # Section heading
    heading = doc.add_heading('Technical Analysis', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    if content_options.get('include_equipment', True):
        doc.add_paragraph("Equipment Specifications:", style='Heading 2')
        
        from app.utils.site_backend import load_site_stage_result
        
        for site in sites_list:
            if "📊 Entire Portfolio" not in site_selection and site.get('name') not in site_selection:
                continue
            
            site_name = site.get('name', 'Unknown')
            
            # Get latest results
            for stage in ['detailed', 'preliminary', 'concept', 'screening']:
                result = load_site_stage_result(site_name, stage)
                if result and result.get('complete'):
                    doc.add_paragraph(f"{site_name} - {stage.capitalize()} Stage:", style='Heading 3')
                    
                    equipment = result.get('equipment', {})
                    if equipment:
                        for eq_type, capacity in equipment.items():
                            doc.add_paragraph(
                                f"  • {eq_type.capitalize()}: {capacity:.1f} MW",
                                style='List Bullet'
                            )
                    else:
                        doc.add_paragraph("  No equipment data available")
                    
                    break


def add_site_information(doc: Document, site_selection: List[str], sites_list: List[Dict]):
    """Add site information section"""
    
    # Section heading
    heading = doc.add_heading('Site Information', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
    
    for site in sites_list:
        if "📊 Entire Portfolio" not in site_selection and site.get('name') not in site_selection:
            continue
        
        site_name = site.get('name', 'Unknown')
        doc.add_paragraph(site_name, style='Heading 2')
        
        # Site details
        details = [
            ("Location", site.get('location', 'N/A')),
            ("IT Capacity", f"{site.get('it_capacity_mw', 0)} MW"),
            ("Facility Capacity", f"{site.get('facility_mw', 0):.0f} MW"),
            ("Land Area", f"{site.get('land_acres', 0)} acres"),
            ("Coordinates", site.get('coordinates', 'N/A'))
        ]
        
        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").font.bold = True
            p.add_run(str(value))
        
        doc.add_paragraph()
